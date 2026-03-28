"""Local file content extraction — parse .md, .txt, .docx, .doc, .html into structured sections."""

import pathlib
import re
from datetime import datetime
from typing import List, Optional, Tuple, TYPE_CHECKING

from .exceptions import LocalFileError
from .models import ArticleInfo, ArticleSection

if TYPE_CHECKING:
    from .config import LocalFilesConfig


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def extract_local_file(
    file_path: str,
    config: "LocalFilesConfig",
) -> Tuple[ArticleInfo, List[ArticleSection]]:
    """Extract content from a local file.  Dispatches by file extension.

    Returns ``(ArticleInfo, sections)``.
    Raises ``LocalFileError`` on failure.
    """
    p = pathlib.Path(file_path)
    if not p.exists():
        raise LocalFileError(f"File not found: {file_path}")

    suffix = p.suffix.lower()
    extractors = {
        ".md": _extract_markdown,
        ".txt": _extract_txt,
        ".docx": _extract_docx,
        ".doc": _extract_doc,
        ".html": _extract_html,
        ".htm": _extract_html,
    }
    extractor = extractors.get(suffix)
    if not extractor:
        raise LocalFileError(f"Unsupported file format: {suffix}")

    return extractor(p, config)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _guess_title(sections: List[ArticleSection]) -> str:
    """Derive title from the first heading, if any."""
    for sec in sections:
        if sec.heading:
            return sec.heading
    return ""


def _file_date(p: pathlib.Path) -> str:
    """Return file modification date as YYYY-MM-DD."""
    return datetime.fromtimestamp(p.stat().st_mtime).strftime("%Y-%m-%d")


def _build_info(
    path: pathlib.Path,
    sections: List[ArticleSection],
    *,
    title: str = "",
    author: str = "",
    publish_date: str = "",
    language: Optional[str] = None,
    description: str = "",
) -> ArticleInfo:
    """Build an ``ArticleInfo`` from file metadata and parsed sections."""
    body = "\n\n".join(
        (sec.heading + "\n" + sec.body if sec.heading else sec.body)
        for sec in sections
    )
    word_count = len(body.split())
    return ArticleInfo(
        title=title or _guess_title(sections) or path.stem,
        url=path.resolve().as_uri(),
        author=author,
        site_name="",
        publish_date=publish_date or _file_date(path),
        language=language,
        description=description,
        word_count=word_count,
        sections=sections,
    )


# ---------------------------------------------------------------------------
# .md  — Markdown files
# ---------------------------------------------------------------------------

def _extract_markdown(
    path: pathlib.Path,
    config: "LocalFilesConfig",
) -> Tuple[ArticleInfo, List[ArticleSection]]:
    text = path.read_text(encoding="utf-8", errors="replace")

    # Parse YAML frontmatter (between --- markers)
    frontmatter: dict = {}
    body = text
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            try:
                import yaml
                frontmatter = yaml.safe_load(parts[1]) or {}
            except Exception:
                pass
            body = parts[2]

    # Reuse the heading parser from pdf.py (skip PDF-specific cleaning)
    from .pdf import parse_markdown_to_sections
    sections = parse_markdown_to_sections(body, pdf_cleanup=False)

    if not sections:
        sections = [ArticleSection(heading="", level=2, body=body.strip())]

    info = _build_info(
        path, sections,
        title=str(frontmatter.get("title", "")),
        author=str(frontmatter.get("author", "")),
        publish_date=str(frontmatter.get("date", "")),
        language=frontmatter.get("language") or frontmatter.get("lang"),
        description=str(frontmatter.get("description", "")),
    )
    return info, sections


# ---------------------------------------------------------------------------
# .txt  — Plain text files
# ---------------------------------------------------------------------------

_RE_ALLCAPS = re.compile(r"^[A-Z][A-Z0-9 :,.\-]{2,}$")
_RE_COLON_HEADING = re.compile(r"^(.{3,60}):\s*$")


def _detect_txt_headings(text: str) -> List[ArticleSection]:
    """Split text into sections by detecting pseudo-headings."""
    lines = text.split("\n")
    sections: List[ArticleSection] = []
    current_heading = ""
    current_lines: List[str] = []

    def _flush():
        body = "\n".join(current_lines).strip()
        if body or current_heading:
            sections.append(ArticleSection(
                heading=current_heading,
                level=2,
                body=body,
            ))

    for line in lines:
        stripped = line.strip()
        is_heading = False
        if stripped and _RE_ALLCAPS.match(stripped):
            is_heading = True
        elif stripped and _RE_COLON_HEADING.match(stripped):
            is_heading = True

        if is_heading:
            _flush()
            current_heading = stripped.rstrip(":")
            current_lines = []
        else:
            current_lines.append(line)

    _flush()
    return sections


def _extract_txt(
    path: pathlib.Path,
    config: "LocalFilesConfig",
) -> Tuple[ArticleInfo, List[ArticleSection]]:
    text = path.read_text(encoding="utf-8", errors="replace")

    if config.detect_txt_headings:
        sections = _detect_txt_headings(text)
    else:
        sections = []

    # Fallback: single section split on double newlines as paragraphs
    if not sections:
        body = re.sub(r"\n{3,}", "\n\n", text.strip())
        sections = [ArticleSection(heading="", level=2, body=body)]

    info = _build_info(path, sections)
    return info, sections


# ---------------------------------------------------------------------------
# .docx  — Word documents (python-docx)
# ---------------------------------------------------------------------------

def _extract_docx(
    path: pathlib.Path,
    config: "LocalFilesConfig",
) -> Tuple[ArticleInfo, List[ArticleSection]]:
    from .deps import ensure_python_docx
    ensure_python_docx()
    import docx

    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        raise LocalFileError(f"Failed to open .docx file: {exc}") from exc

    sections: List[ArticleSection] = []
    current_heading = ""
    current_level = 2
    current_paragraphs: List[str] = []

    def _flush():
        nonlocal current_heading, current_paragraphs
        body = "\n\n".join(current_paragraphs)
        if body or current_heading:
            sections.append(ArticleSection(
                heading=current_heading,
                level=current_level,
                body=body,
            ))
        current_heading = ""
        current_paragraphs = []

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading"):
            _flush()
            current_heading = para.text.strip()
            level_match = re.search(r"(\d+)", style_name)
            current_level = int(level_match.group(1)) if level_match else 2
            current_paragraphs = []
        else:
            text = para.text.strip()
            if text:
                current_paragraphs.append(text)

    _flush()

    # Optionally include tables
    if config.include_tables and doc.tables:
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                # Markdown table with header separator
                table_md = rows[0] + "\n" + " | ".join(["---"] * len(table.rows[0].cells))
                if len(rows) > 1:
                    table_md += "\n" + "\n".join(rows[1:])
                sections.append(ArticleSection(
                    heading="",
                    level=2,
                    body=table_md,
                ))

    if not sections:
        raise LocalFileError("No content found in .docx file.")

    # Extract metadata from document properties
    props = doc.core_properties
    title = props.title or ""
    author = props.author or ""
    date_val = props.created or props.modified
    publish_date = date_val.strftime("%Y-%m-%d") if date_val else ""

    info = _build_info(
        path, sections,
        title=title,
        author=author,
        publish_date=publish_date,
    )
    return info, sections


# ---------------------------------------------------------------------------
# .doc  — Legacy Word (mammoth → HTML → trafilatura)
# ---------------------------------------------------------------------------

def _extract_doc(
    path: pathlib.Path,
    config: "LocalFilesConfig",
) -> Tuple[ArticleInfo, List[ArticleSection]]:
    from .deps import ensure_mammoth
    ensure_mammoth()
    import mammoth

    try:
        with open(path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html = result.value
    except Exception as exc:
        raise LocalFileError(f"Failed to convert .doc file: {exc}") from exc

    if not html or not html.strip():
        raise LocalFileError("No content extracted from .doc file.")

    # Reuse the article extraction pipeline for the converted HTML
    from .article import extract_article
    from .config import ArticlesConfig

    articles_config = ArticlesConfig(
        include_tables=config.include_tables,
        min_content_length=config.min_content_length,
    )
    info, sections = extract_article(html, path.resolve().as_uri(), articles_config)

    # Override with file-based metadata
    if not info.title or info.title == "Untitled":
        info.title = _guess_title(sections) or path.stem
    info.publish_date = info.publish_date if info.publish_date != "unknown" else _file_date(path)

    return info, sections


# ---------------------------------------------------------------------------
# .html / .htm  — Local HTML files (trafilatura)
# ---------------------------------------------------------------------------

def _extract_html(
    path: pathlib.Path,
    config: "LocalFilesConfig",
) -> Tuple[ArticleInfo, List[ArticleSection]]:
    html = path.read_text(encoding="utf-8", errors="replace")

    from .article import extract_article
    from .config import ArticlesConfig

    articles_config = ArticlesConfig(
        include_tables=config.include_tables,
        min_content_length=config.min_content_length,
    )
    info, sections = extract_article(html, path.resolve().as_uri(), articles_config)

    # Override with file-based metadata
    if not info.title or info.title == "Untitled":
        info.title = _guess_title(sections) or path.stem
    info.publish_date = info.publish_date if info.publish_date != "unknown" else _file_date(path)

    return info, sections
